"""
Document Ingestion Pipeline

Supports multiple file formats:
- Plain text
- Markdown
- HTML
- PDF
- DOCX
- JSON
- CSV

Extracts and normalizes:
- Title
- Content
- Language
- Author
- Source
- URL
- Tags
- Categories
- Metadata
- Created date
- Updated date
"""

import asyncio
import hashlib
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from uuid import uuid4

import aiofiles

logger = logging.getLogger("nebula.search.ingestion")


class DocumentStatus(str, Enum):
    """Document processing status"""
    PENDING = "pending"
    PROCESSING = "processing"
    INDEXED = "indexed"
    FAILED = "failed"
    REJECTED = "rejected"


class FileType(str, Enum):
    """Supported file types"""
    TEXT = "text"
    MARKDOWN = "markdown"
    HTML = "html"
    PDF = "pdf"
    DOCX = "docx"
    JSON = "json"
    CSV = "csv"
    UNKNOWN = "unknown"


@dataclass
class DocumentMetadata:
    """Document metadata"""
    title: str = ""
    author: str = ""
    source: str = ""
    url: str = ""
    tags: List[str] = field(default_factory=list)
    categories: List[str] = field(default_factory=list)
    created_date: Optional[datetime] = None
    updated_date: Optional[datetime] = None
    language: str = "en"
    custom_metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Document:
    """Ingested document"""
    id: str
    filename: str
    file_type: FileType
    content: str
    metadata: DocumentMetadata
    status: DocumentStatus
    file_size: int
    checksum: str
    created_at: datetime
    updated_at: datetime
    error_message: Optional[str] = None
    processing_time_ms: Optional[float] = None


class DocumentIngester:
    """Document ingestion pipeline with multi-format support"""

    def __init__(self):
        self.supported_extensions = {
            '.txt': FileType.TEXT,
            '.md': FileType.MARKDOWN,
            '.markdown': FileType.MARKDOWN,
            '.html': FileType.HTML,
            '.htm': FileType.HTML,
            '.pdf': FileType.PDF,
            '.docx': FileType.DOCX,
            '.json': FileType.JSON,
            '.csv': FileType.CSV,
        }
        self.max_file_size = 50 * 1024 * 1024  # 50MB

    def detect_file_type(self, filename: str) -> FileType:
        """Detect file type from extension"""
        ext = Path(filename).suffix.lower()
        return self.supported_extensions.get(ext, FileType.UNKNOWN)

    async def ingest_file(
        self,
        file_path: Union[str, Path],
        metadata: Optional[DocumentMetadata] = None,
    ) -> Document:
        """
        Ingest a single file.

        Args:
            file_path: Path to the file
            metadata: Optional metadata (if None, will be extracted)

        Returns:
            Document object
        """
        file_path = Path(file_path)
        start_time = datetime.utcnow()

        try:
            # Validate file exists
            if not file_path.exists():
                raise ValueError(f"File not found: {file_path}")

            # Check file size
            file_size = file_path.stat().st_size
            if file_size > self.max_file_size:
                raise ValueError(f"File too large: {file_size} bytes (max {self.max_file_size})")

            # Calculate checksum
            checksum = await self._calculate_checksum(file_path)

            # Detect file type
            file_type = self.detect_file_type(file_path.name)

            # Extract metadata if not provided
            if metadata is None:
                metadata = await self._extract_metadata(file_path, file_type)

            # Parse content
            content = await self._parse_file(file_path, file_type)

            # Validate content
            if not content or len(content.strip()) < 10:
                raise ValueError("File content too short or empty")

            # Sanitize content
            content = self._sanitize_content(content)

            processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000

            return Document(
                id=str(uuid4()),
                filename=file_path.name,
                file_type=file_type,
                content=content,
                metadata=metadata,
                status=DocumentStatus.INDEXED,
                file_size=file_size,
                checksum=checksum,
                created_at=start_time,
                updated_at=datetime.utcnow(),
                processing_time_ms=processing_time,
            )

        except Exception as e:
            logger.error(f"Failed to ingest {file_path}: {e}")
            processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            return Document(
                id=str(uuid4()),
                filename=file_path.name,
                file_type=self.detect_file_type(file_path.name),
                content="",
                metadata=metadata or DocumentMetadata(),
                status=DocumentStatus.FAILED,
                file_size=0,
                checksum="",
                created_at=start_time,
                updated_at=datetime.utcnow(),
                error_message=str(e),
                processing_time_ms=processing_time,
            )

    async def ingest_content(
        self,
        content: str,
        filename: str = "content.txt",
        metadata: Optional[DocumentMetadata] = None,
    ) -> Document:
        """
        Ingest content directly from string.

        Args:
            content: Text content
            filename: Filename (for type detection)
            metadata: Optional metadata

        Returns:
            Document object
        """
        start_time = datetime.utcnow()
        file_type = self.detect_file_type(filename)

        try:
            # Validate content
            if not content or len(content.strip()) < 10:
                raise ValueError("Content too short or empty")

            # Sanitize content
            content = self._sanitize_content(content)

            # Extract metadata if not provided
            if metadata is None:
                metadata = self._extract_metadata_from_content(content, filename)

            processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            checksum = hashlib.sha256(content.encode()).hexdigest()

            return Document(
                id=str(uuid4()),
                filename=filename,
                file_type=file_type,
                content=content,
                metadata=metadata,
                status=DocumentStatus.INDEXED,
                file_size=len(content.encode()),
                checksum=checksum,
                created_at=start_time,
                updated_at=datetime.utcnow(),
                processing_time_ms=processing_time,
            )

        except Exception as e:
            logger.error(f"Failed to ingest content: {e}")
            processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            return Document(
                id=str(uuid4()),
                filename=filename,
                file_type=file_type,
                content="",
                metadata=metadata or DocumentMetadata(),
                status=DocumentStatus.FAILED,
                file_size=0,
                checksum="",
                created_at=start_time,
                updated_at=datetime.utcnow(),
                error_message=str(e),
                processing_time_ms=processing_time,
            )

    async def _parse_file(self, file_path: Path, file_type: FileType) -> str:
        """Parse file content based on type"""
        if file_type == FileType.TEXT:
            return await self._parse_text(file_path)
        elif file_type == FileType.MARKDOWN:
            return await self._parse_markdown(file_path)
        elif file_type == FileType.HTML:
            return await self._parse_html(file_path)
        elif file_type == FileType.PDF:
            return await self._parse_pdf(file_path)
        elif file_type == FileType.DOCX:
            return await self._parse_docx(file_path)
        elif file_type == FileType.JSON:
            return await self._parse_json(file_path)
        elif file_type == FileType.CSV:
            return await self._parse_csv(file_path)
        else:
            return await self._parse_text(file_path)

    async def _parse_text(self, file_path: Path) -> str:
        """Parse plain text file"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            return await f.read()

    async def _parse_markdown(self, file_path: Path) -> str:
        """Parse markdown file"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        # Remove markdown syntax (basic implementation)
        content = re.sub(r'#+\s+', '', content)
        content = re.sub(r'\*{1,2}([^\*]+)\*{1,2}', r'\1', content)
        content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
        content = re.sub(r'```.*?```', '', content, flags=re.DOTALL)
        return content.strip()

    async def _parse_html(self, file_path: Path) -> str:
        """Parse HTML file"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        # Remove HTML tags (basic implementation)
        content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL)
        content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL)
        content = re.sub(r'<[^>]+>', ' ', content)
        content = re.sub(r'\s+', ' ', content)
        return content.strip()

    async def _parse_pdf(self, file_path: Path) -> str:
        """Parse PDF file"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                return text.strip()
        except ImportError:
            logger.warning("PyPDF2 not installed, returning empty content for PDF")
            return ""
        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            raise

    async def _parse_docx(self, file_path: Path) -> str:
        """Parse DOCX file"""
        try:
            import docx
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except ImportError:
            logger.warning("python-docx not installed, returning empty content for DOCX")
            return ""
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise

    async def _parse_json(self, file_path: Path) -> str:
        """Parse JSON file"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        import json
        data = json.loads(content)
        return json.dumps(data, indent=2)

    async def _parse_csv(self, file_path: Path) -> str:
        """Parse CSV file"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        lines = content.strip().split('\n')
        return '\n'.join(lines)

    async def _extract_metadata(
        self, file_path: Path, file_type: FileType
    ) -> DocumentMetadata:
        """Extract metadata from file"""
        metadata = DocumentMetadata(
            title=file_path.stem,
            source=str(file_path.parent),
        )

        # Set created/updated dates from file stats
        stat = file_path.stat()
        metadata.created_date = datetime.fromtimestamp(stat.st_ctime)
        metadata.updated_date = datetime.fromtimestamp(stat.st_mtime)

        # Try to extract from content
        try:
            content = await self._parse_file(file_path, file_type)
            metadata = self._extract_metadata_from_content(content, file_path.name)
        except Exception as e:
            logger.warning(f"Failed to extract metadata from content: {e}")

        return metadata

    def _extract_metadata_from_content(
        self, content: str, filename: str
    ) -> DocumentMetadata:
        """Extract metadata from content"""
        metadata = DocumentMetadata(
            title=Path(filename).stem,
            source="",
        )

        # Try to extract title from first line
        lines = content.strip().split('\n')
        if lines:
            first_line = lines[0].strip()
            if len(first_line) < 200:
                metadata.title = first_line

        # Extract dates (basic pattern matching)
        date_patterns = [
            r'\d{4}-\d{2}-\d{2}',
            r'\d{2}/\d{2}/\d{4}',
            r'\d{2}-\d{2}-\d{4}',
        ]
        for pattern in date_patterns:
            matches = re.findall(pattern, content[:1000])
            if matches:
                try:
                    metadata.created_date = datetime.strptime(matches[0], '%Y-%m-%d')
                    break
                except ValueError:
                    continue

        # Extract author (look for "by Author" or "Author:")
        author_patterns = [
            r'by\s+([A-Z][a-z]+ [A-Z][a-z]+)',
            r'author[:\s]+([A-Z][a-z]+ [A-Z][a-z]+)',
        ]
        for pattern in author_patterns:
            match = re.search(pattern, content[:2000], re.IGNORECASE)
            if match:
                metadata.author = match.group(1)
                break

        return metadata

    async def _calculate_checksum(self, file_path: Path) -> str:
        """Calculate SHA256 checksum of file"""
        sha256_hash = hashlib.sha256()
        async with aiofiles.open(file_path, 'rb') as f:
            while chunk := await f.read(8192):
                sha256_hash.update(chunk)
        return sha256_hash.hexdigest()

    def _sanitize_content(self, content: str) -> str:
        """Sanitize content"""
        # Remove control characters
        content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', content)
        # Normalize whitespace
        content = re.sub(r'\s+', ' ', content)
        # Strip leading/trailing whitespace
        content = content.strip()
        return content


# Global instance
document_ingester = DocumentIngester()